from __future__ import annotations

import copy
import re
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\Allotrope\PDD\Active methodology")
TEMPLATE = Path(r"D:\Allotrope\PDD\VCS templates\VCS-Project-Description-Template-v4.4-FINAL2.docx")
YAML_PATH = Path(r"D:\Allotrope\PDD\intake.inegol_acm0022_from_pd.yaml")
OUT = ROOT / "generated_pdd" / "INEGOL_VCS_Project_Description_v4.4_draft.docx"
ISSUE_DATE = date(2026, 5, 7)

MISSING = "[TBD - source evidence/calculation required]"


def strip_comment(line: str) -> str:
    in_quote = False
    quote = ""
    for i, ch in enumerate(line):
        if ch in ("'", '"'):
            if not in_quote:
                in_quote = True
                quote = ch
            elif quote == ch:
                in_quote = False
        if ch == "#" and not in_quote:
            return line[:i]
    return line


def parse_scalar(value: str):
    value = value.strip()
    if value == "":
        return ""
    if value in ("null", "Null", "NULL", "~"):
        return None
    if value in ("true", "True", "TRUE"):
        return True
    if value in ("false", "False", "FALSE"):
        return False
    if value.startswith("[") and value.endswith("]"):
        inner = value[1:-1].strip()
        if not inner:
            return []
        parts = []
        buf = ""
        in_quote = False
        quote = ""
        for ch in inner:
            if ch in ("'", '"'):
                if not in_quote:
                    in_quote = True
                    quote = ch
                elif quote == ch:
                    in_quote = False
                buf += ch
            elif ch == "," and not in_quote:
                parts.append(parse_scalar(buf.strip()))
                buf = ""
            else:
                buf += ch
        if buf.strip():
            parts.append(parse_scalar(buf.strip()))
        return parts
    if (value.startswith('"') and value.endswith('"')) or (value.startswith("'") and value.endswith("'")):
        return value[1:-1]
    if re.fullmatch(r"-?\d+", value):
        return int(value)
    if re.fullmatch(r"-?\d+\.\d+", value):
        return float(value)
    return value


def load_simple_yaml(path: Path):
    raw_lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    lines = []
    for raw in raw_lines:
        raw = strip_comment(raw).rstrip()
        if not raw.strip():
            continue
        indent = len(raw) - len(raw.lstrip(" "))
        lines.append((indent, raw.strip()))

    def parse_block(i: int, indent: int):
        if i >= len(lines):
            return {}, i
        is_list = lines[i][1].startswith("- ")
        if is_list:
            arr = []
            while i < len(lines):
                ind, text = lines[i]
                if ind < indent:
                    break
                if ind != indent or not text.startswith("- "):
                    break
                item_text = text[2:].strip()
                if not item_text:
                    child, i = parse_block(i + 1, indent + 2)
                    arr.append(child)
                    continue
                if ":" in item_text and not item_text.startswith(("'", '"')):
                    key, val = item_text.split(":", 1)
                    item = {}
                    if val.strip():
                        item[key.strip()] = parse_scalar(val.strip())
                        i += 1
                    else:
                        child, i = parse_block(i + 1, indent + 2)
                        item[key.strip()] = child
                    while i < len(lines):
                        ni, nt = lines[i]
                        if ni <= indent:
                            break
                        if ni == indent + 2 and not nt.startswith("- "):
                            if ":" not in nt:
                                break
                            k, v = nt.split(":", 1)
                            if v.strip():
                                item[k.strip()] = parse_scalar(v.strip())
                                i += 1
                            else:
                                child, i = parse_block(i + 1, ni + 2)
                                item[k.strip()] = child
                        else:
                            break
                    arr.append(item)
                else:
                    arr.append(parse_scalar(item_text))
                    i += 1
            return arr, i

        obj = {}
        while i < len(lines):
            ind, text = lines[i]
            if ind < indent:
                break
            if ind != indent:
                break
            if ":" not in text:
                i += 1
                continue
            key, val = text.split(":", 1)
            key = key.strip()
            if val.strip():
                obj[key] = parse_scalar(val.strip())
                i += 1
            else:
                child, i = parse_block(i + 1, indent + 2)
                obj[key] = child
        return obj, i

    parsed, _ = parse_block(0, 0)
    return parsed


def repair_text(value):
    if isinstance(value, dict):
        return {k: repair_text(v) for k, v in value.items()}
    if isinstance(value, list):
        return [repair_text(v) for v in value]
    if not isinstance(value, str):
        return value
    repaired = value
    try:
        candidate = value.encode("latin1").decode("utf-8")
        bad_before = sum(value.count(x) for x in ("Ã", "Ä", "Å"))
        bad_after = sum(candidate.count(x) for x in ("Ã", "Ä", "Å"))
        if bad_after < bad_before:
            repaired = candidate
    except UnicodeError:
        pass
    return repaired


def fmt_date(value: str | None) -> str:
    if not value:
        return MISSING
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").strftime("%d-%B-%Y")
    except ValueError:
        return str(value)


def parse_date(value: str) -> date:
    return datetime.strptime(str(value), "%Y-%m-%d").date()


def fmt_num(value, decimals=0) -> str:
    if value is None:
        return MISSING
    if isinstance(value, str):
        return value
    if decimals == 0:
        return f"{value:,.0f}"
    return f"{value:,.{decimals}f}"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, margin_twips=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, font_size=9):
    cell.text = ""
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margin(cell)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def add_table(doc, rows, widths=None, header=True, font_size=8.7):
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        set_row_cant_split(table.rows[r_idx])
        if header and r_idx == 0:
            set_row_repeat_header(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            is_head = header and r_idx == 0
            set_cell_text(cell, value, bold=is_head, font_size=font_size)
            if is_head:
                set_cell_shading(cell, "D9EAF7")
            elif c_idx == 0 and len(row) == 2:
                set_cell_shading(cell, "F2F6FA")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_para(doc, text="", style=None, bold_label=None):
    if style and style not in [s.name for s in doc.styles]:
        style = None
    p = doc.add_paragraph(style=style)
    if bold_label:
        run = p.add_run(bold_label)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    return p


def add_bullets(doc, items):
    style = "Bullets"
    if style not in [s.name for s in doc.styles]:
        style = None
    for item in items:
        if item:
            add_para(doc, item, style=style)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level > 1 else 16)
    p.paragraph_format.space_after = Pt(5)
    return p


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_page_break(doc):
    doc.add_page_break()


def vintage_periods(start: date, end: date):
    rows = []
    current = start
    while current <= end:
        year_end = date(current.year, 12, 31)
        period_end = min(year_end, end)
        rows.append((current, period_end))
        current = period_end + timedelta(days=1)
    return rows


def toc_rows():
    return [
        ["1", "Project Details"],
        ["1.1", "Summary Description of the Project"],
        ["1.2", "Audit History"],
        ["1.3", "Sectoral Scope and Project Type"],
        ["1.4", "Project Eligibility"],
        ["1.5", "Project Design"],
        ["1.6", "Project Proponent"],
        ["1.7", "Other Entities Involved in the Project"],
        ["1.8", "Ownership"],
        ["1.9", "Project Start Date"],
        ["1.10", "Project Crediting Period"],
        ["1.11", "Project Scale and Estimated GHG Emission Reductions or Removals"],
        ["1.12", "Description of the Project Activity"],
        ["1.13", "Project Location"],
        ["1.14", "Conditions Prior to Project Initiation"],
        ["1.15", "Compliance with Laws, Statutes and Other Regulatory Frameworks"],
        ["1.16", "Double Counting and Participation under Other GHG Programs"],
        ["1.17", "Double Claiming, Other Forms of Credit, and Scope 3 Emissions"],
        ["1.18", "Sustainable Development Contributions"],
        ["1.19", "Additional Information Relevant to the Project"],
        ["2", "Safeguards and Stakeholder Engagement"],
        ["2.1", "Stakeholder Engagement and Consultation"],
        ["2.2", "Risks to Stakeholders and the Environment"],
        ["2.3", "Respect for Human Rights and Equity"],
        ["2.4", "Ecosystem Health"],
        ["3", "Application of Methodology"],
        ["3.1", "Title and Reference of Methodology"],
        ["3.2", "Applicability of Methodology"],
        ["3.3", "Project Boundary"],
        ["3.4", "Baseline Scenario"],
        ["3.5", "Additionality"],
        ["3.6", "Methodology Deviations"],
        ["4", "Quantification of Estimated GHG Emission Reductions and Removals"],
        ["4.1", "Baseline Emissions"],
        ["4.2", "Project Emissions"],
        ["4.3", "Leakage Emissions"],
        ["4.4", "Estimated GHG Emission Reductions and Carbon Dioxide Removals"],
        ["5", "Monitoring"],
        ["5.1", "Data and Parameters Available at Validation"],
        ["5.2", "Data and Parameters Monitored"],
        ["5.3", "Monitoring Plan"],
        ["Appendix 1", "Commercially Sensitive Information"],
        ["Appendix 2", "Data Gaps and Assumptions Register"],
    ]


def build_document():
    data = repair_text(load_simple_yaml(YAML_PATH))
    ident = data["project_identity"]
    parties = data["parties"]
    tech = data["technical_design"]
    loc = data["location"]
    meth = data["methodology"]
    proponents = parties.get("project_proponents", [])
    other_entities = parties.get("other_entities", [])
    feedstock_desc = tech.get("feedstock_or_waste_type", MISSING)
    facility_manager = other_entities[0].get("org_name", MISSING) if other_entities else MISSING

    title = ident["project_title"]
    start = parse_date(ident["crediting_period"]["start_date"])
    end = parse_date(ident["crediting_period"]["end_date"])
    prepared_by = ident.get("prepared_by", MISSING)

    doc = Document(str(TEMPLATE))
    clear_body(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.orientation = WD_ORIENT.PORTRAIT

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    add_para(doc, "VCS Project Description - Template v4.4 draft", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Generated from intake YAML and supplied methodology/template reference files.", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER

    cover_rows = [
        ["Project title", title],
        ["Project ID", ident.get("project_id", MISSING)],
        ["Crediting period", f"{fmt_date(ident['crediting_period']['start_date'])} to {fmt_date(ident['crediting_period']['end_date'])}"],
        ["Original date of issue", fmt_date(ident.get("registry_listing_date"))],
        ["Most recent date of issue", ISSUE_DATE.strftime("%d-%B-%Y")],
        ["Version", "Draft 0.1"],
        ["VCS Standard Version", ident.get("vcs_standard_version", MISSING)],
        ["Prepared by", prepared_by],
    ]
    add_table(doc, cover_rows, widths=[Inches(2.2), Inches(4.8)], header=False, font_size=9.3)

    add_para(
        doc,
        "This draft uses project-specific facts from the intake file. Items not present in the source intake are marked as TBD and should be completed from the validated calculation spreadsheet, permits, stakeholder records, and supporting evidence package before submission.",
    )
    add_page_break(doc)

    add_heading(doc, "Contents", 1)
    add_table(doc, [["Section", "Title"]] + toc_rows(), widths=[Inches(1.0), Inches(5.8)], header=True, font_size=8.3)
    add_page_break(doc)

    add_heading(doc, "Project Details", 1)
    add_heading(doc, "Summary Description of the Project", 2)
    add_para(
        doc,
        f"The project activity is the {title}, an integrated municipal solid waste management and biogas-to-electricity project located in {loc['city_or_district']}, {loc['region_state']}, {loc['country']}. The source feedstock is described as: {feedstock_desc} Recyclable fractions are recovered where applicable, and residual fractions are managed through the waste management system.",
    )
    add_para(
        doc,
        f"The installed and operational electricity generation capacity is {fmt_num(tech.get('installed_capacity_mw'), 3)} MWe, consisting of six gas engines commissioned between {fmt_date(tech['gas_engine_commissioning'][0]['commissioning_date'])} and {fmt_date(tech['gas_engine_commissioning'][-1]['commissioning_date'])}. The generation license allows a total capacity of {fmt_num(tech.get('license_capacity_mw'), 3)} MWe. Expected net electricity generation is {fmt_num(tech.get('expected_net_generation_mwh_per_year'), 3)} MWh per year and {fmt_num(tech.get('expected_net_generation_mwh_total_crediting_period'), 3)} MWh over the first crediting period.",
    )
    add_para(
        doc,
        f"The project reduces greenhouse gas emissions by diverting fresh municipal solid waste from disposal at a solid waste disposal site, avoiding methane emissions from anaerobic decay in the baseline, and generating renewable electricity from recovered biogas that is exported to the Turkish national grid via the {loc.get('grid_connection_point', MISSING)}. The project applies ACM0022, version {meth.get('methodology_version')}, for alternative waste treatment processes.",
    )

    add_heading(doc, "Audit History", 2)
    audit_rows = [["Audit type", "Period", "Program", "Validation/verification body name", "Years"]]
    for audit in ident.get("audit_history", []):
        audit_rows.append([
            audit.get("audit_type", MISSING).title(),
            audit.get("period", MISSING),
            audit.get("program", MISSING),
            audit.get("vvb_name", MISSING),
            str(audit.get("number_of_years", MISSING)),
        ])
    add_table(doc, audit_rows, widths=[Inches(1.1), Inches(1.7), Inches(0.8), Inches(1.7), Inches(0.6)], header=True)

    add_heading(doc, "Sectoral Scope and Project Type", 2)
    add_table(
        doc,
        [
            ["Sectoral scope", "\n".join(tech.get("sectoral_scopes", []))],
            ["Project activity type", tech.get("project_activity_type", MISSING)],
        ],
        widths=[Inches(1.8), Inches(5.3)],
        header=False,
    )

    add_heading(doc, "Project Eligibility", 2)
    add_heading(doc, "General eligibility", 3)
    add_para(
        doc,
        f"The project is a non-AFOLU waste handling and disposal activity within the scope of the VCS Program. The intake identifies the project as {ident.get('pipeline_listing_status', MISSING)} on the Verra pipeline on {fmt_date(ident.get('registry_listing_date'))}. The validation opening meeting date is recorded as {fmt_date(ident.get('validation_opening_date'))}. Validation deadline exemptions are recorded as {', '.join(fmt_date(x) for x in ident.get('validation_deadline_exemptions', []))}.",
    )
    add_para(
        doc,
        "The project activity is not identified as excluded under the VCS Standard in the intake. Final eligibility should be confirmed against the applicable VCS Standard version and the evidence package submitted to the VVB.",
    )
    add_heading(doc, "AFOLU project eligibility", 3)
    add_para(doc, "Not applicable. The project is a non-AFOLU municipal solid waste management and energy generation project.")
    add_heading(doc, "Transfer project eligibility", 3)
    add_para(doc, "No transfer from another GHG program is identified in the intake. If any prior GHG program registration exists, the no-double-issuance evidence must be added before submission.")

    add_heading(doc, "Project Design", 2)
    add_para(doc, f"The project is designed as a {ident.get('project_design_type', 'single_location').replace('_', ' ')} project. The intake identifies grouped project eligibility as {ident.get('grouped_project_eligibility', MISSING).replace('_', ' ')}.")
    add_heading(doc, "Grouped project design", 3)
    add_para(doc, "Not applicable. No grouped project instances or inclusion criteria are proposed in the source intake.")

    add_heading(doc, "Project Proponent", 2)
    for prop in proponents:
        add_table(
            doc,
            [
                ["Organization name", prop.get("org_name", MISSING)],
                ["Contact person", prop.get("contact_name", MISSING)],
                ["Title", prop.get("title", MISSING)],
                ["Address", prop.get("address", MISSING)],
                ["Telephone", prop.get("telephone", MISSING)],
                ["Email", prop.get("email", MISSING)],
            ],
            widths=[Inches(1.7), Inches(5.4)],
            header=False,
        )

    add_heading(doc, "Other Entities Involved in the Project", 2)
    for ent in other_entities:
        add_table(
            doc,
            [
                ["Organization name", ent.get("org_name", MISSING)],
                ["Role in the project", ent.get("role", MISSING)],
                ["Contact person", ent.get("contact_name", MISSING)],
                ["Title", ent.get("title", MISSING)],
                ["Address", ent.get("address", MISSING)],
                ["Telephone", ent.get("telephone", MISSING)],
                ["Email", ent.get("email", MISSING)],
            ],
            widths=[Inches(1.7), Inches(5.4)],
            header=False,
        )

    add_heading(doc, "Ownership", 2)
    add_para(doc, parties.get("ownership_summary", MISSING))
    add_para(doc, f"Ownership evidence identified in the intake: {', '.join(parties.get('ownership_evidence_ids', [])) or MISSING}.")

    add_heading(doc, "Project Start Date", 2)
    add_table(
        doc,
        [
            ["Project start date", fmt_date(ident.get("project_start_date"))],
            ["Justification", ident.get("project_start_date_justification", MISSING)],
        ],
        widths=[Inches(1.7), Inches(5.4)],
        header=False,
    )

    add_heading(doc, "Project Crediting Period", 2)
    cp = ident["crediting_period"]
    add_table(
        doc,
        [
            ["Crediting period", "Seven years, twice renewable"],
            ["Start and end date of first crediting period", f"{fmt_date(cp.get('start_date'))} to {fmt_date(cp.get('end_date'))}"],
        ],
        widths=[Inches(2.3), Inches(4.8)],
        header=False,
    )

    add_heading(doc, "Project Scale and Estimated GHG Emission Reductions or Removals", 2)
    add_para(
        doc,
        f"The intake identifies the project scale class as {ident.get('project_scale_class', MISSING)} tCO2e per year. The source project scale label is recorded as {ident.get('source_project_scale_label', MISSING)}; this should be reconciled with the final emission reduction calculation spreadsheet before submission.",
    )
    scale_rows = [["Calendar year of crediting period", "Estimated GHG emission reductions or removals (tCO2e)"]]
    for a, b in vintage_periods(start, end):
        scale_rows.append([f"{a.strftime('%d-%B-%Y')} to {b.strftime('%d-%B-%Y')}", MISSING])
    scale_rows.append(["Total", MISSING])
    add_table(doc, scale_rows, widths=[Inches(3.1), Inches(3.0)], header=True)

    add_heading(doc, "Description of the Project Activity", 2)
    add_para(doc, tech.get("equipment_summary", MISSING))
    add_para(
        doc,
        f"The expected municipal solid waste throughput is {fmt_num(tech.get('expected_waste_throughput_tonnes_per_year'), 2)} tonnes/year, equivalent to approximately {fmt_num(tech.get('expected_waste_throughput_tonnes_per_day'), 2)} tonnes/day. The intake states that approximately {fmt_num(tech.get('biomethanization_suitable_fraction') * 100 if tech.get('biomethanization_suitable_fraction') is not None else None, 1)}% of the waste stream is suitable for biomethanization, corresponding to {fmt_num(tech.get('biomethanization_suitable_waste_tonnes_per_day'), 1)} tonnes/day.",
    )
    rdf = tech.get("rdf_capacity", {})
    add_para(
        doc,
        f"The RDF preparation system has a maximum capacity of {fmt_num(rdf.get('maximum_capacity_tonnes_per_hour'), 0)} tonnes/hour. Planned RDF throughput is stated as {fmt_num(rdf.get('planned_capacity_tonnes_per_day_2024'), 0)} tonnes/day in 2024 and {fmt_num(rdf.get('planned_capacity_tonnes_per_day_2035'), 0)} tonnes/day in 2035.",
    )
    add_para(doc, tech.get("implementation_schedule", MISSING))
    add_table(
        doc,
        [["Engine", "Model", "Commissioning date"]]
        + [[e.get("engine", ""), e.get("model", ""), fmt_date(e.get("commissioning_date"))] for e in tech.get("gas_engine_commissioning", [])],
        widths=[Inches(0.8), Inches(3.8), Inches(1.5)],
        header=True,
    )

    add_heading(doc, "Project Location", 2)
    add_para(doc, f"The project site is located at {loc.get('site_description', MISSING)} The site area is {fmt_num(loc.get('site_area_m2'), 2)} m2.")
    add_table(
        doc,
        [
            ["Country", loc.get("country", MISSING)],
            ["Region/state", loc.get("region_state", MISSING)],
            ["City/district", loc.get("city_or_district", MISSING)],
            ["Grid connection point", loc.get("grid_connection_point", MISSING)],
            ["Boundary evidence", ", ".join(loc.get("boundary_kml_evidence_ids", [])) or MISSING],
        ],
        widths=[Inches(1.8), Inches(5.3)],
        header=False,
    )
    coord_rows = [["Point", "Latitude", "Longitude"]]
    for i, c in enumerate(loc.get("coordinates", []), start=1):
        coord_rows.append([str(i), f"{c.get('lat'):.7f}", f"{c.get('lon'):.7f}"])
    add_table(doc, coord_rows, widths=[Inches(0.7), Inches(2.0), Inches(2.0)], header=True)

    add_heading(doc, "Conditions Prior to Project Initiation", 2)
    add_para(doc, tech.get("pre_project_waste_management", MISSING))
    add_para(doc, "The baseline service levels include disposal of municipal solid waste in the existing waste disposal system and electricity supplied by grid-connected electricity generation plants. The project activity provides alternative waste treatment, biogas recovery and use, and renewable electricity generation.")

    add_heading(doc, "Compliance with Laws, Statutes and Other Regulatory Frameworks", 2)
    add_para(doc, f"The intake records an EMRA generation license dated 17-December-2020 and identifies {facility_manager} as holding the legal rights of the project. Final submission should include the generation license, environmental permits, waste management permits/municipal authorizations, grid connection documentation, and any EIA or EIA-exemption evidence applicable to the facility.")

    add_heading(doc, "Double Counting and Participation under Other GHG Programs", 2)
    add_heading(doc, "No Double Issuance", 3)
    add_para(doc, f"No other GHG program registration is identified in the intake. Evidence to support no double issuance remains {MISSING}.")
    add_heading(doc, "Registration in Other GHG Programs", 3)
    add_para(doc, "No registration under another GHG program is identified in the intake.")
    add_heading(doc, "Projects Rejected by Other GHG Programs", 3)
    add_para(doc, "No rejection by another GHG program is identified in the intake.")

    add_heading(doc, "Double Claiming, Other Forms of Credit, and Scope 3 Emissions", 2)
    add_heading(doc, "No Double Claiming with Emissions Trading Programs or Binding Emission Limits", 3)
    add_para(doc, f"No emissions trading program or binding emission limit interaction is identified in the intake. Final evidence is {MISSING}.")
    add_heading(doc, "No Double Claiming with Other Forms of Environmental Credit", 3)
    add_para(doc, f"No other environmental crediting for the same GHG benefit is identified in the intake. Final evidence is {MISSING}.")
    add_heading(doc, "Supply Chain (Scope 3) Emissions", 3)
    add_para(doc, "The project reduces methane emissions from municipal solid waste management and displaces grid electricity. Any claims by downstream users of RDF, recycled materials, or electricity attributes should be reviewed and documented to prevent double claiming.")

    add_heading(doc, "Sustainable Development Contributions", 2)
    add_para(doc, "The project contributes to sustainable development by improving municipal waste handling, recovering recyclable and combustible fractions, treating organic waste through biomethanization, recovering biogas for renewable electricity generation, reducing methane emissions, and supporting local operational employment.")
    add_table(
        doc,
        [
            ["Sustainable development area", "Project contribution", "Monitoring approach"],
            ["Climate mitigation", "Avoided methane from fresh municipal waste and displacement of grid electricity.", "Annual emission reduction calculations and electricity export/import records."],
            ["Clean energy", "Renewable electricity from recovered biogas is exported to the Turkish national grid.", "Metered electricity supplied to and consumed from the grid."],
            ["Responsible waste management", "Mechanical separation, biomethanization, RDF preparation, and residual waste management improve integrated municipal waste treatment.", "Waste reception records, process records, RDF production records, and residual disposal records."],
            ["Local employment and skills", "The facility requires operational, maintenance, monitoring, and health and safety staff.", "Employment and training records, where available."],
        ],
        widths=[Inches(1.7), Inches(3.0), Inches(2.4)],
        header=True,
    )

    add_heading(doc, "Additional Information Relevant to the Project", 2)
    add_heading(doc, "Leakage Management", 3)
    add_para(doc, ident.get("additional_information", MISSING))
    add_heading(doc, "Commercially Sensitive Information", 3)
    add_para(doc, "No commercially sensitive information is intentionally excluded from this draft except where source evidence is marked as TBD. Any final exclusions should be listed in Appendix 1 with justification.")
    add_heading(doc, "Further Information", 3)
    add_para(doc, "The final project description should be cross-checked against the validated calculation spreadsheet, permits, stakeholder consultation records, grid connection documents, equipment specifications, and ownership documentation.")

    add_page_break(doc)
    add_heading(doc, "Safeguards and Stakeholder Engagement", 1)
    add_heading(doc, "Stakeholder Engagement and Consultation", 2)
    add_heading(doc, "Stakeholder Identification", 3)
    add_table(
        doc,
        [
            ["Stakeholder Identification", f"Stakeholders expected to be relevant include Bursa Metropolitan Municipality and district municipalities, nearby residents and local administrative units around {loc.get('site_description', 'the project site')}, facility workers and contractors, waste collection/transport entities, grid/operator interfaces, recyclable/RDF off-takers, and relevant public authorities. Site-specific stakeholder lists and evidence are TBD."],
            ["Legal or customary tenure/access rights", f"The project is located adjacent to the existing solid waste landfill facility on a defined site area of {fmt_num(loc.get('site_area_m2'), 2)} m2. Evidence of land rights, municipal access rights, and any customary access considerations is {MISSING}."],
            ["Stakeholder diversity and changes over time", "Stakeholders may include public authorities, local communities, facility workers, waste sector contractors, and commercial off-takers. Demographic and vulnerability information is TBD."],
            ["Expected changes in well-being", "Expected effects include improved waste management, reduced landfill methane potential, employment, and possible localized traffic, odour, noise, and occupational health and safety risks that require management."],
            ["Location of stakeholders", f"Primary stakeholders are expected to be located in and around the project site and the service area described in the intake: {feedstock_desc}"],
            ["Location of resources", "Relevant resources include the project site, municipal waste streams, access roads, grid connection infrastructure, and surrounding community resources. Site-specific mapping evidence is TBD."],
        ],
        widths=[Inches(2.1), Inches(5.0)],
        header=False,
        font_size=8.2,
    )
    add_heading(doc, "Stakeholder Consultation and Ongoing Communication", 3)
    add_table(
        doc,
        [
            ["Date of stakeholder consultation", MISSING],
            ["Stakeholder engagement process", "Stakeholder consultation records were not included in the intake. The final PD should describe meeting announcements, participant groups, language/accessibility arrangements, meeting minutes, attendance records, and how input was documented."],
            ["Consultation outcome", MISSING],
            ["Ongoing communication", "The final PD should provide contact points, grievance channels, response timelines, and records management procedures for ongoing communication with stakeholders."],
            ["Stakeholder input", "No site-specific input was provided in the intake. Final design updates or justification for no updates should be added after reviewing consultation evidence."],
        ],
        widths=[Inches(2.1), Inches(5.0)],
        header=False,
        font_size=8.2,
    )
    add_heading(doc, "Free Prior and Informed Consent", 3)
    add_table(
        doc,
        [
            ["Obtaining consent", "No Indigenous Peoples, local communities with customary tenure rights, or FPIC-triggering circumstances are identified in the intake. This conclusion should be confirmed with stakeholder mapping and land/access rights evidence."],
            ["Outcome of FPIC", "Not applicable based on the intake, subject to confirmation through stakeholder identification and land/access rights review."],
        ],
        widths=[Inches(2.1), Inches(5.0)],
        header=False,
        font_size=8.2,
    )
    add_heading(doc, "Grievance Redress Procedure", 3)
    add_table(
        doc,
        [
            ["Development process", MISSING],
            ["Grievance redress procedure", "The final PD should include the procedure for receiving, recording, investigating, responding to, and closing grievances, including responsible persons, response timelines, appeal/escalation process, and records retention."],
        ],
        widths=[Inches(2.1), Inches(5.0)],
        header=False,
        font_size=8.2,
    )
    add_heading(doc, "Public Comments", 3)
    add_table(
        doc,
        [["Comments received", "Actions taken"], [MISSING, MISSING]],
        widths=[Inches(3.3), Inches(3.3)],
        header=True,
        font_size=8.2,
    )

    add_heading(doc, "Risks to Stakeholders and the Environment", 2)
    add_heading(doc, "Management Experience", 3)
    add_para(doc, "The project proponent group and facility manager are identified in the intake. Final submission should include evidence of experience in waste management, power generation, occupational health and safety, environmental compliance, and VCS project implementation.")
    add_heading(doc, "Risk Assessment", 3)
    add_table(
        doc,
        [
            ["Risk category", "Risks identified", "Mitigation or preventative measure(s) taken"],
            ["Natural and human-induced risks to stakeholders' wellbeing", "Odour, traffic, noise, fire/explosion risk from biogas, and emergency events.", "Operate gas collection, booster, flare, safety valves, pressure sensors, emergency procedures, traffic controls, and community communication channels."],
            ["Risks to stakeholder participation", "Low access to information or limited awareness of consultation channels.", "Provide clear public contact points, accessible consultation notices, and grievance channels."],
            ["Environmental impacts", "Leachate, wastewater, residual waste, air emissions, and odour.", "Operate leachate handling, process controls, permitted discharge/treatment systems, regular inspections, and environmental monitoring."],
            ["Occupational health and safety", "Mechanical equipment, confined spaces, biogas, electrical equipment, and vehicle movements.", "Use training, PPE, lockout/tagout, gas detection, permit-to-work procedures, and emergency drills."],
            ["Project continuity", "Equipment downtime and flare/engine interruptions.", "Preventive maintenance, emergency flaring, spare parts management, and monitoring of engine and gas system performance."],
        ],
        widths=[Inches(1.6), Inches(2.5), Inches(3.0)],
        header=True,
        font_size=8.0,
    )

    add_heading(doc, "Respect for Human Rights and Equity", 2)
    add_heading(doc, "Labor and Work", 3)
    add_table(
        doc,
        [
            ["Risk category", "Risks identified", "Mitigation or preventative measure(s) taken"],
            ["Discrimination", "No specific risk identified in the intake.", "Apply equal opportunity employment practices and grievance channels."],
            ["Sexual harassment", "No specific risk identified in the intake.", "Apply workplace conduct policies, confidential reporting, and disciplinary procedures."],
            ["Forced labor", "No specific risk identified in the intake.", "Comply with Turkish labor law and supplier/contractor requirements."],
            ["Child labor", "No specific risk identified in the intake.", "Verify worker age and contractor compliance."],
            ["Freedom of association", "No specific risk identified in the intake.", "Comply with applicable labor rights and maintain worker communication channels."],
            ["Occupational health and safety", "Biogas, mechanical, traffic, and electrical hazards.", "Training, PPE, inspections, emergency response, and permit-to-work controls."],
        ],
        widths=[Inches(1.6), Inches(2.5), Inches(3.0)],
        header=True,
        font_size=8.0,
    )
    add_heading(doc, "Human Rights", 3)
    add_para(doc, "No human rights risk specific to the project is identified in the intake. Final submission should document the screening process and applicable mitigation measures.")
    add_heading(doc, "Indigenous Peoples and Cultural Heritage", 3)
    add_para(doc, "No Indigenous Peoples or cultural heritage impacts are identified in the intake. This should be confirmed with stakeholder mapping and permitting evidence.")
    add_heading(doc, "Property Rights", 3)
    add_para(doc, "No displacement or adverse property-right impact is identified in the intake. Land tenure, access rights, and municipal authorization evidence should be attached or referenced.")
    add_heading(doc, "Benefit Sharing", 3)
    add_para(doc, "No separate benefit-sharing plan is identified in the intake. The project benefits are expected to arise through improved municipal waste management, employment, renewable electricity generation, and climate mitigation. If affected stakeholder groups require a benefit-sharing plan, it should be added.")

    add_heading(doc, "Ecosystem Health", 2)
    add_para(doc, "The project is a waste management and energy project at/adjacent to an existing solid waste facility, not an AFOLU land-use project. The final PD should confirm site ecological screening, permitting, and any required environmental management measures.")
    add_table(
        doc,
        [
            ["Risk category", "Risks identified", "Mitigation or preventative measure(s) taken"],
            ["Impacts on biodiversity and ecosystems", "No site-specific biodiversity risk identified in the intake.", "Confirm through permit/EIA screening and operate within permitted project footprint."],
            ["Soil degradation and soil erosion", "Potential localized risk during construction/maintenance.", "Use erosion controls, paved/managed operational areas, and stormwater/leachate management."],
            ["Water consumption and stress", "Water use and wastewater/leachate handling require control.", "Monitor water use, collect/treat leachate and wastewater according to permits."],
            ["Pollution", "Potential odour, air emissions, leachate, and noise.", "Apply process controls, biogas management, emergency flare, and environmental monitoring."],
        ],
        widths=[Inches(1.6), Inches(2.5), Inches(3.0)],
        header=True,
        font_size=8.0,
    )
    add_heading(doc, "Rare, Threatened, and Endangered Species", 3)
    add_para(doc, f"No rare, threatened, or endangered species information was provided in the intake. Site-specific screening is {MISSING}.")
    add_heading(doc, "Introduction of Species", 3)
    add_para(doc, "Not applicable. The project does not involve planting, species introduction, or monoculture establishment.")
    add_heading(doc, "Ecosystem Conversion", 3)
    add_para(doc, "Not applicable as an AFOLU conversion requirement; nevertheless, final permitting evidence should confirm that the project footprint is authorized and does not convert protected natural ecosystems.")

    add_page_break(doc)
    add_heading(doc, "Application of Methodology", 1)
    add_heading(doc, "Title and Reference of Methodology", 2)
    method_rows = [["Type", "Reference ID", "Title", "Version"]]
    method_rows.append(["Methodology", meth.get("methodology_id", "ACM0022"), "Alternative waste treatment processes", meth.get("methodology_version", "03.0")])
    for tool in meth.get("tools", []):
        method_rows.append(["Tool", tool.get("id", ""), tool.get("title", ""), tool.get("version", "")])
    add_table(doc, method_rows, widths=[Inches(0.9), Inches(1.0), Inches(4.2), Inches(0.8)], header=True, font_size=7.7)

    add_heading(doc, "Applicability of Methodology", 2)
    app_rows = [
        ["Methodology/tool", "Applicability condition", "Justification of compliance"],
        ["ACM0022", "Project installs and operates new plant(s) for treatment of fresh waste through eligible processes such as anaerobic digestion with biogas recovery and/or mechanical/thermal treatment to produce RDF/SB.", "The project includes mechanical separation, biomethanization, biogas recovery/use in gas engines, emergency flaring, and RDF preparation for municipal solid waste."],
        ["ACM0022", "Fresh waste would otherwise be disposed in a solid waste disposal site, with or without partial LFG capture.", tech.get("pre_project_waste_management", MISSING)],
        ["ACM0022", "Fresh waste and products from the project plant are not stored on-site under anaerobic conditions outside the eligible treatment process.", "The intake describes direct biogas routing to engines, emergency flare availability, gas balloon storage, bypass lines, pressure sensors, and safety valves. Final operating procedures should confirm storage conditions."],
        ["ACM0022", "Wastewater discharge from the project activity is treated according to applicable regulations.", f"Leachate handling is included in the facility description. Permit and discharge evidence is {MISSING}."],
        ["ACM0022", "The project activity does not reduce recycling that would occur in the absence of the project.", f"The mechanical separation facility recovers recyclable fractions where applicable. Baseline recycling evidence is {MISSING}."],
        ["ACM0022", "Hazardous wastes/wastewater are not eligible.", "The feedstock is municipal solid waste from the named districts. Final waste acceptance procedures should confirm exclusion of hazardous waste."],
        ["TOOL02", "Baseline scenario identification and additionality demonstration are required.", "The project uses the combined tool; final baseline selection and additionality evidence are to be completed with investment/common practice analysis."],
        ["TOOL14", "Anaerobic digester project and leakage emissions are addressed where anaerobic digestion is used.", "The project uses biomethanization and biogas recovery; monitored biogas and operational data will support calculations."],
        ["TOOL05/TOOL07", "Electricity consumption/generation and grid emission factor are addressed where electricity is consumed or exported.", "The project exports renewable electricity to the Turkish grid and consumes auxiliary electricity; metering data and applicable grid EF are required."],
    ]
    add_table(doc, app_rows, widths=[Inches(1.0), Inches(3.1), Inches(3.0)], header=True, font_size=7.6)

    add_heading(doc, "Project Boundary", 2)
    add_para(doc, "The project boundary includes the solid waste disposal site baseline for waste that would otherwise be disposed, the project alternative waste treatment facilities, on-site biogas management and electricity generation equipment, on-site fuel and electricity use, wastewater/leachate handling associated with the treatment process, and the Turkish grid plants connected to the electricity system displaced by project electricity exports. Waste collection and transport are excluded from the boundary under ACM0022.")
    add_table(
        doc,
        [
            ["Boundary element", "Included project facilities/processes"],
            ["Waste reception and mechanical separation", f"{feedstock_desc}; sorting/separation of organic, recyclable, combustible, and residual fractions."],
            ["Biomethanization", "Anaerobic treatment of organic fractions, biogas recovery, gas handling, storage/safety devices, and associated digestate/leachate handling."],
            ["Electricity generation", "Six gas engines totaling 8.484 MWe, booster unit, metering, and connection to the Turkish national grid."],
            ["Emergency flaring", "Emergency flare for destruction of excess biogas when engines cannot consume all produced gas."],
            ["RDF preparation", "Mechanical/thermal preparation of combustible fractions for RDF, with residual waste management."],
            ["Excluded", "Waste collection and transport outside the facility boundary, unless required by a specific leakage calculation."],
        ],
        widths=[Inches(2.1), Inches(5.0)],
        header=True,
        font_size=8.0,
    )
    add_table(
        doc,
        [
            ["Scenario", "Source", "Gas", "Included?", "Justification"],
            ["Baseline", "SWDS disposal of fresh waste", "CH4", "Yes", "Major methane source avoided by alternative waste treatment."],
            ["Baseline", "SWDS disposal of fresh waste", "CO2", "No", "Biogenic CO2 from decomposition of fresh waste is not accounted under ACM0022."],
            ["Baseline", "SWDS disposal of fresh waste", "N2O", "No", "Excluded as small compared with CH4 and conservative."],
            ["Baseline", "Electricity generation displaced by project exports", "CO2", "Yes", "Project electricity exported to the grid displaces grid-connected generation."],
            ["Baseline", "Electricity generation", "CH4/N2O", "No", "Excluded for simplification/conservativeness under the methodology."],
            ["Project", "On-site fossil fuel consumption", "CO2", "Yes", "Included where fossil fuel is consumed by project equipment."],
            ["Project", "On-site electricity use", "CO2", "Yes", "Auxiliary electricity consumption may be an important project emission source."],
            ["Project", "Anaerobic digestion and gas handling", "CH4", "Yes", "Potential methane leakage from digesters/gas systems and incomplete combustion."],
            ["Project", "Flaring", "CH4", "Yes", "Incomplete combustion from emergency flaring is addressed using the relevant tool."],
            ["Project", "RDF/SB production and use", "CO2/CH4/N2O", "Yes/No as applicable", "Included where fossil-based fractions or off-site end use create emissions required by ACM0022; final treatment and end-use evidence is required."],
            ["Leakage", "Anaerobic digestion and RDF/SB by-products/end use", "CH4/CO2/N2O", "Yes as applicable", "Leakage sources are treated according to ACM0022 and relevant tools."],
        ],
        widths=[Inches(0.8), Inches(2.2), Inches(0.7), Inches(0.8), Inches(2.6)],
        header=True,
        font_size=7.2,
    )

    add_heading(doc, "Baseline Scenario", 2)
    add_para(doc, "The baseline scenario is selected using ACM0022 and the combined tool to identify the baseline scenario and demonstrate additionality. Based on the intake, the baseline for fresh municipal solid waste is disposal in a SWDS with partial landfill gas capture and flaring. For electricity generated by the project activity, the baseline is electricity generation in existing and/or new grid-connected electricity plants.")
    add_bullets(
        doc,
        [
            "Waste baseline: fresh municipal solid waste disposed at a SWDS with partial LFG capture and flaring.",
            "Energy baseline: electricity supplied by grid-connected electricity plants serving the Turkish national grid.",
            "Project alternative: alternative treatment of municipal solid waste through mechanical separation, biomethanization with biogas recovery and use, RDF preparation, and renewable electricity generation.",
        ],
    )

    add_heading(doc, "Additionality", 2)
    add_heading(doc, "Regulatory Surplus", 3)
    add_para(doc, "The intake does not identify a regulation mandating the project activity at a compliance rate that would make the activity non-additional. Final regulatory surplus evidence should include applicable Turkish waste, renewable energy, landfill gas, environmental, and licensing requirements and demonstrate that the project activity is not legally required or is eligible under ACM0022 where compliance rates are below the applicable threshold.")
    add_heading(doc, "Additionality Methods", 3)
    add_para(doc, "Additionality is to be demonstrated using the combined tool to identify the baseline scenario and demonstrate additionality, supported by the investment analysis tool and common practice tool listed in the intake. The final PD should include the selected analysis route, all inputs and sources, sensitivity analysis, and common practice assessment sufficient for a reader to reproduce the conclusion.")
    add_table(
        doc,
        [
            ["Step", "Draft application", "Evidence status"],
            ["Identify realistic and credible alternatives", "Alternatives include continued disposal in the SWDS with partial LFG capture/flaring, implementation of alternative waste treatment without carbon revenue, and project implementation with VCS revenue.", MISSING],
            ["Investment analysis", "Apply TOOL27 to assess financial attractiveness of the project activity without carbon revenue.", MISSING],
            ["Barrier analysis, if selected", "Not selected in this draft; to be confirmed.", MISSING],
            ["Common practice", f"Apply TOOL24 to assess similar integrated MSW/biogas/RDF projects in {loc.get('country', 'the host country')}.", MISSING],
            ["Conclusion", "The project is expected to be additional subject to completion of investment and common practice evidence.", MISSING],
        ],
        widths=[Inches(1.3), Inches(3.5), Inches(2.2)],
        header=True,
        font_size=8.0,
    )

    add_heading(doc, "Methodology Deviations", 2)
    add_para(doc, "No methodology deviations are identified in the intake. If any deviation is applied during validation, it must be described and justified here, including evidence that it does not reduce conservativeness and relates only to monitoring or measurement criteria/procedures.")

    add_page_break(doc)
    add_heading(doc, "Quantification of Estimated GHG Emission Reductions and Removals", 1)
    add_heading(doc, "Baseline Emissions", 2)
    add_para(doc, "Baseline emissions are quantified in accordance with ACM0022. For this project, relevant baseline components include methane emissions from fresh waste that would have been disposed in the baseline SWDS and baseline emissions from electricity generation displaced by renewable electricity exported to the grid.")
    add_bullets(doc, [
        "BE_y = sum of applicable baseline components for each treatment process and year, adjusted for any applicable regulatory compliance discount factor.",
        "BE_CH4,SWDS,y is calculated using TOOL04 for the amount and composition of fresh waste prevented from disposal.",
        "BE_EC,y is calculated using the applicable electricity generation/consumption tool and Turkish grid emission factor for net electricity supplied by the project.",
    ])
    add_heading(doc, "Project Emissions", 2)
    add_para(doc, "Project emissions include applicable emissions from anaerobic digestion, gas handling, flaring, on-site fossil fuel use, on-site electricity consumption, RDF/SB production and end-use treatment where required, and any other sources included by ACM0022 and the referenced tools.")
    add_bullets(doc, [
        "PE_AD,y is determined using the tool for project and leakage emissions from anaerobic digesters.",
        "PE_flare,y is determined using the tool for project emissions from flaring where emergency flaring occurs.",
        "PE_EC,y accounts for auxiliary electricity consumption from the grid.",
        "PE_FC,y accounts for fossil fuel consumption by project equipment, if any.",
    ])
    add_heading(doc, "Leakage Emissions", 2)
    add_para(doc, "Leakage emissions are assessed for anaerobic digestion and RDF/SB pathways according to ACM0022. The intake identifies leakage management measures including direct biogas routing to engines, emergency flaring, pressure sensors, bypass lines, HDPE pipes, gas balloon storage, and safety valves. Final leakage calculations require monitored or estimated data on digestate/by-products, RDF/SB end use, residual disposal, and any off-site anaerobic decomposition risks.")
    add_heading(doc, "Estimated GHG Emission Reductions and Carbon Dioxide Removals", 2)
    add_para(doc, "Emission reductions are calculated as baseline emissions minus project emissions and leakage emissions. Carbon dioxide removals are not claimed by this non-AFOLU waste management project.")
    add_para(doc, f"The intake provides expected electricity generation but does not provide annual baseline, project, leakage, or net emission reduction values. The table below is therefore prepared with required vintage periods and marked {MISSING} pending the calculation spreadsheet.")
    er_rows = [["Vintage period", "Estimated baseline emissions (tCO2e)", "Estimated project emissions (tCO2e)", "Estimated leakage emissions (tCO2e)", "Estimated reductions (tCO2e)", "Estimated removals (tCO2e)", "Estimated VCUs (tCO2e)"]]
    for a, b in vintage_periods(start, end):
        er_rows.append([f"{a.strftime('%d-%b-%Y')} to {b.strftime('%d-%b-%Y')}", MISSING, MISSING, MISSING, MISSING, "0", MISSING])
    er_rows.append(["Total", MISSING, MISSING, MISSING, MISSING, "0", MISSING])
    add_table(doc, er_rows, widths=[Inches(1.3), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.8), Inches(1.0)], header=True, font_size=6.8)

    add_page_break(doc)
    add_heading(doc, "Monitoring", 1)
    add_heading(doc, "Data and Parameters Available at Validation", 2)
    fixed_params = [
        ["Installed_capacity", "MW", "Installed operational electricity generation capacity", fmt_num(tech.get("installed_capacity_mw"), 3), "Project technical/equipment records in intake", "Capacity should be verified against equipment commissioning records and license."],
        ["License_capacity", "MW", "Licensed total electricity generation capacity", fmt_num(tech.get("license_capacity_mw"), 3), "EMRA generation license referenced in intake", "License evidence to be attached."],
        ["Expected_waste_throughput", "tonnes/year", "Expected municipal solid waste throughput", fmt_num(tech.get("expected_waste_throughput_tonnes_per_year"), 2), "Source PD/intake", "Used for ex-ante estimates; monitored values supersede during verification."],
        ["Biomethanization_suitable_fraction", "fraction", "Expected fraction suitable for biomethanization", fmt_num(tech.get("biomethanization_suitable_fraction"), 2), "Source PD/intake", "Waste characterization evidence required."],
        ["Project_start_date", "date", "Project start date", fmt_date(ident.get("project_start_date")), "Partial acceptance date for first gas engine", "Used for crediting period and vintage boundaries."],
        ["Grid_connection_point", "text", "Grid connection point", loc.get("grid_connection_point", MISSING), "Source PD/intake", "Grid connection documentation required."],
    ]
    add_table(
        doc,
        [["Data / Parameter", "Unit", "Description", "Value applied", "Source", "Comments"]] + fixed_params,
        widths=[Inches(1.3), Inches(0.8), Inches(2.0), Inches(1.0), Inches(1.2), Inches(1.5)],
        header=True,
        font_size=7.0,
    )

    add_heading(doc, "Data and Parameters Monitored", 2)
    monitored_params = [
        ["W_MSW,y", "tonnes", "Municipal solid waste received/treated by the project", "Continuous/weighbridge or operational records; aggregated monthly and annually", "Weighbridge/waste acceptance system", "Calibration and reconciliation with municipal records"],
        ["Waste fractions", "fraction or %", "Composition of waste by relevant ACM0022 categories", "Sampling/characterization according to approved procedure", "Waste characterization records/lab analysis", "Sampling plan and QA/QC required"],
        ["EG_export,y", "MWh", "Electricity supplied to the grid", "Continuous meter reading; monthly and annual aggregation", "Revenue/export meter", "Meter calibration, cross-check with grid invoices"],
        ["EG_import,y", "MWh", "Electricity consumed from the grid", "Continuous meter reading; monthly and annual aggregation", "Import meter/invoices", "Used for project emissions from electricity consumption"],
        ["Q_biogas,y", "Nm3 or m3", "Biogas flow to engines/flare", "Continuous flow measurement; aggregated monthly and annually", "Biogas flow meters/SCADA", "Meter calibration and data completeness checks"],
        ["CH4_biogas,y", "% v/v", "Methane content of biogas", "Continuous analyzer or periodic measurement per monitoring plan", "Gas analyzer/lab records", "QA/QC and calibration required"],
        ["Flare operation", "hours/flow/temperature", "Emergency flare operation and destruction conditions", "Event-based and continuous where applicable", "Flare logs/SCADA", "Demonstrate combustion and destruction efficiency"],
        ["Fossil fuel use", "litres, kg, or GJ", "Fossil fuels consumed by project equipment", "As consumed; monthly/annual aggregation", "Fuel purchase/use records", "Apply TOOL03 where applicable"],
        ["RDF production/end use", "tonnes", "RDF produced and sent to end users or otherwise managed", "Operational records; monthly/annual aggregation", "RDF production/offtake records", "Needed for RDF/SB leakage/end-use assessment"],
        ["Residual waste/by-products", "tonnes", "Residuals, digestate, leachate or by-products sent to disposal/treatment/use", "Operational records; monthly/annual aggregation", "Facility logs/manifests", "Needed for leakage and project emissions"],
    ]
    add_table(
        doc,
        [["Data / Parameter", "Unit", "Description", "Frequency", "Monitoring equipment/source", "QA/QC procedures"]] + monitored_params,
        widths=[Inches(1.1), Inches(0.8), Inches(1.9), Inches(1.2), Inches(1.3), Inches(1.5)],
        header=True,
        font_size=6.8,
    )

    add_heading(doc, "Monitoring Plan", 2)
    add_para(doc, "Monitoring will be implemented by trained facility personnel under the responsibility of the project proponent/facility manager. Data will be collected from calibrated meters, SCADA/plant records, weighbridge records, waste characterization studies, invoices, laboratory analyses, operating logs, and maintenance/calibration records.")
    add_bullets(
        doc,
        [
            "Waste reception data will be recorded through weighbridge or equivalent systems and reconciled with municipal/operational records.",
            "Electricity exported to and imported from the grid will be recorded using revenue-grade meters and reconciled with grid operator or settlement records.",
            "Biogas flow, methane content, engine operation, and flare events will be recorded through meters, analyzers, SCADA, and operating logs.",
            "RDF production, residual waste, digestate, leachate, and by-product movements will be recorded using facility logs, weighbridge records, manifests, and off-take/disposal documents.",
            "Monitoring equipment will be calibrated according to manufacturer recommendations, applicable standards, and the validated monitoring plan.",
            "Data will be archived electronically and/or physically for the period required by the VCS Program and made available to the VVB during verification.",
            "Missing data, meter failures, non-conformances, and corrective actions will be documented using conservative procedures consistent with the methodology and monitoring plan.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Appendix 1: Commercially Sensitive Information", 1)
    add_table(
        doc,
        [["Section", "Information", "Justification"], ["None identified in this draft", "No commercially sensitive information has been intentionally excluded.", "N/A"]],
        widths=[Inches(1.5), Inches(3.5), Inches(2.0)],
        header=True,
        font_size=8.0,
    )

    add_heading(doc, "Appendix 2: Data Gaps and Assumptions Register", 1)
    add_table(
        doc,
        [
            ["Topic", "Gap / assumption", "Needed evidence"],
            ["Emission reductions", "Annual baseline, project, leakage, and net ER values are not provided in the intake.", "Validated emission reduction calculation spreadsheet."],
            ["Project scale", "Intake contains a scale class and a source scale label that may not be aligned.", "Final ER estimates and VCS scale classification check."],
            ["Stakeholder consultation", "No project-specific consultation date, attendance, comments, or outcomes provided.", "Consultation plan, notices, attendance, minutes, comments, and responses."],
            ["Permits and legal compliance", "EMRA license is referenced, but full permit set is not included.", "Generation license, environmental permits, waste management authorizations, grid connection evidence, EIA/EIA exemption if applicable."],
            ["Additionality", "Investment/common practice inputs and conclusion are not provided.", "TOOL02/TOOL27/TOOL24 analysis, inputs, sources, and sensitivity tests."],
            ["Waste characterization", "Biomethanization-suitable fraction is provided, but detailed composition support is not included.", "Waste characterization study and sampling procedure."],
            ["Monitoring equipment", "Meter/analyzer model, serial number, accuracy, and calibration details are not provided.", "Equipment inventory and calibration certificates."],
        ],
        widths=[Inches(1.5), Inches(3.3), Inches(2.3)],
        header=True,
        font_size=8.0,
    )

    doc.core_properties.title = title
    doc.core_properties.subject = "VCS Project Description draft"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "VCS, ACM0022, project description, Inegol"
    doc.core_properties.created = datetime.combine(ISSUE_DATE, datetime.min.time())
    doc.core_properties.modified = datetime.combine(ISSUE_DATE, datetime.min.time())
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    output = build_document()
    print(output)
