"""DOCX export — produces a structured Word document from a DraftRun.

Uses python-docx to generate a Verra VCS PDD-formatted document with:
- Per-section headings matching canonical schema numbering
- DraftSection text inserted under each heading
- REVIEW REQUIRED blocks styled in yellow where confidence is LOW/UNSUPPORTED
- A review status table as an appendix
- Provenance citations inline
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import structlog
import yaml

logger = structlog.get_logger()

_DRAFT_RUNS_DIR = Path(__file__).parent.parent.parent.parent / "data" / "runs"
_SCHEMA_PATH = Path(__file__).parent.parent.parent.parent / "schemas" / "pdd_section_schema.yaml"


def export_run_to_docx(
    run_id: str,
    output_path: Path | None = None,
    project_name: str = "",
) -> Path | None:
    """Export a DraftRun to a formatted DOCX file.

    Args:
        run_id: The run identifier (without .json extension).
        output_path: Optional explicit output path. If None, uses data/runs/{run_id}.docx.
        project_name: Project name for the document title page.

    Returns:
        Path to the generated DOCX file, or None if the run was not found.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.error("python-docx not installed — cannot export DOCX. Run: pip install python-docx")
        return None

    import json

    run_path = _DRAFT_RUNS_DIR / f"{run_id}.json"
    if not run_path.exists():
        logger.error("draft_run_not_found", run_id=run_id, path=str(run_path))
        return None

    with open(run_path, encoding="utf-8") as f:
        run_data = json.load(f)

    doc = Document()

    _add_title_page(doc, project_name or run_data.get("project_name", "Unknown Project"), run_id)
    _add_meta_table(doc, run_data)

    doc.add_heading("Project Design Document", level=1)

    sections = run_data.get("sections", [])
    schema = _load_schema()

    for sec_def in schema.get("sections", []):
        sid = sec_def["section_id"]
        sec_heading = sec_def.get("canonical_heading", f"Section {sid}")
        doc.add_heading(sec_heading, level=1)

        for ss_def in sec_def.get("sub_sections", []):
            ssid = ss_def["sub_section_id"]
            draft = next(
                (s for s in sections if s["section_id"] == sid and s["sub_section_id"] == ssid),
                None,
            )
            ss_heading = ss_def.get("heading", f"{sid}.{ssid}")
            doc.add_heading(ss_heading, level=2)

            if draft:
                text = draft.get("text", "")
                confidence = draft.get("confidence", "UNKNOWN")
                issues = draft.get("issues", [])

                if draft.get("provenance"):
                    prov_para = doc.add_paragraph()
                    prov_para.add_run("Provenance: ").bold = True
                    prov_para.add_run("; ".join(draft.get("provenance", [])))
                    prov_para.style = "Quote"

                if text:
                    para = doc.add_paragraph(text)
                    if confidence in ("LOW", "UNSUPPORTED"):
                        _highlight_paragraph(para, RGBColor(0xFF, 0xFF, 0x00))
                else:
                    doc.add_paragraph("[No content drafted yet]").style = "Intense Quote"

                if issues:
                    issues_para = doc.add_paragraph()
                    issues_para.add_run("Review Issues:").bold = True
                    for issue in issues:
                        run = issues_para.add_paragraph(f"  • {issue}")
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                doc.add_paragraph("[Not drafted]").style = "Intense Quote"

    _add_review_status_table(doc, sections)
    _add_page_numbers(doc)

    if output_path is None:
        output_path = _DRAFT_RUNS_DIR / f"{run_id}.docx"
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("docx_exported", run_id=run_id, path=str(output_path))
    return output_path


def _add_title_page(doc: Any, project_name: str, run_id: str) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Design Document")
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Verra VCS — Waste-to-Energy")
    run2.font.size = Pt(16)

    doc.add_paragraph()
    proj_para = doc.add_paragraph()
    proj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = proj_para.add_run(project_name)
    run3.bold = True
    run3.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Run ID: {run_id}\n")
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run("Status: DRAFT — Requires Human Review")

    doc.add_page_break()


def _add_meta_table(doc: Any, run_data: dict[str, Any]) -> None:
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    meta = [
        ("Project", run_data.get("project_name", "—")),
        ("Provider", run_data.get("provider", "noop")),
        ("Total Sections", str(len(run_data.get("sections", [])))),
    ]
    for i, (k, v) in enumerate(meta):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()


def _add_review_status_table(doc: Any, sections: list[dict[str, Any]]) -> None:
    doc.add_heading("Section Review Status", level=1)
    table = doc.add_table(rows=len(sections) + 1, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Section", "Sub-section", "Confidence", "Issues"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    for i, sec in enumerate(sections, 1):
        table.rows[i].cells[0].text = sec.get("section_id", "")
        table.rows[i].cells[1].text = sec.get("sub_section_id", "")
        table.rows[i].cells[2].text = sec.get("confidence", "UNKNOWN")
        table.rows[i].cells[3].text = str(len(sec.get("issues", [])))


def _highlight_paragraph(para: Any, color: Any) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for run in para.runs:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "FFFF00")
        run._r.get_or_add_rPr().append(shading)


def _add_page_numbers(doc: Any) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.alignment = 1
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)


def _load_schema() -> dict[str, Any]:
    with open(_SCHEMA_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f)
