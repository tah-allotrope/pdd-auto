"""LLM-based document extraction: arbitrary documents -> ProjectInput YAML.

Accepts DOCX, PDF, or plain text files, extracts text, sends to OpenAI GPT-4
with the extraction prompt, and parses the structured YAML output into a
validated ProjectInput instance.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import structlog
import yaml

from pdd_agent.llm.provider import BaseProvider
from schemas.project_input import ProjectInput, ExtractionProvenance

logger = structlog.get_logger()

_PROMPT_PATH = Path(__file__).parent.parent.parent.parent / "prompts" / "extract_project_input.md"
_MAX_DOC_CHARS = 80_000


class ExtractionError(Exception):
    """Raised when document extraction fails."""


def _read_text_from_file(path: Path) -> str:
    """Extract plain text from DOCX, PDF, or plain text file."""
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx_text(path)
    elif suffix == ".pdf":
        return _extract_pdf_text(path)
    elif suffix in (".txt", ".md", ".text", ".yaml", ".yml"):
        return path.read_text(encoding="utf-8")
    else:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")


def _extract_docx_text(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ExtractionError(
            "python-docx not installed. Install via: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _extract_pdf_text(path: Path) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError("pypdf not installed. Install via: pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _load_prompt_template() -> str:
    """Load the extraction prompt template."""
    if not _PROMPT_PATH.exists():
        raise ExtractionError(f"Extraction prompt not found at {_PROMPT_PATH}")
    return _PROMPT_PATH.read_text(encoding="utf-8")


def _build_extraction_prompt(document_text: str) -> str:
    """Build the full extraction prompt with document text inserted."""
    template = _load_prompt_template()
    truncated = document_text[:_MAX_DOC_CHARS]
    if len(document_text) > _MAX_DOC_CHARS:
        truncated += f"\n\n[... document truncated at {_MAX_DOC_CHARS:,} characters ...]"
    return template.replace("{{DOCUMENT_TEXT}}", truncated)


def _parse_yaml_from_response(response_text: str) -> dict[str, Any]:
    """Parse YAML from LLM response, handling markdown fences."""
    text = response_text.strip()
    fence_match = re.search(r"```(?:yaml)?\s*\n(.*?)```", text, re.DOTALL)
    if fence_match:
        text = fence_match.group(1).strip()

    try:
        parsed = yaml.safe_load(text)
    except yaml.YAMLError as exc:
        raise ExtractionError(f"Failed to parse YAML from LLM response: {exc}") from exc

    if not isinstance(parsed, dict):
        raise ExtractionError(f"Expected YAML dict, got {type(parsed).__name__}")

    if "_error" in parsed:
        raise ExtractionError(
            f"LLM reported extraction error: {parsed['_error']} — {parsed.get('_reason', '')}"
        )

    return parsed


def _clean_missing_markers(data: dict[str, Any]) -> tuple[dict[str, Any], list[str]]:
    """Replace [MISSING] string markers with appropriate defaults, tracking which fields were missing."""
    missing_fields: list[str] = []

    def _walk(obj: Any, path: str = "") -> Any:
        if isinstance(obj, dict):
            return {k: _walk(v, f"{path}.{k}" if path else k) for k, v in obj.items()}
        if isinstance(obj, list):
            return [_walk(item, f"{path}[{i}]") for i, item in enumerate(obj)]
        if isinstance(obj, str) and "[MISSING]" in obj:
            missing_fields.append(path)
            return None
        if isinstance(obj, str) and obj.startswith("[INFERENCE:"):
            return obj
        return obj

    cleaned = _walk(data)
    return cleaned, missing_fields


def _build_provenance(
    raw_data: dict[str, Any],
    missing_fields: list[str],
    source_path: str | None,
    model_name: str | None,
) -> ExtractionProvenance:
    """Build provenance metadata from extraction results."""
    extraction_meta = raw_data.pop("_extraction", {})
    extracted = extraction_meta.get("extracted_fields", [])
    defaulted = extraction_meta.get("defaulted_fields", [])

    return ExtractionProvenance(
        extracted_fields=extracted,
        defaulted_fields=defaulted,
        missing_fields=missing_fields,
        source_document=source_path,
        extraction_model=model_name,
    )


def _apply_schema_defaults(data: dict[str, Any]) -> dict[str, Any]:
    """Apply minimal schema defaults to make the data parseable by ProjectInput."""
    defaults: dict[str, Any] = {
        "methodology_applicability": {"eligibility_checklist": {}},
        "quantification": {},
        "monitoring": {
            "parameters_monitored": [],
            "data_management": "Not specified",
        },
        "safeguards": {"no_net_harm_statement": "Not specified"},
        "compliance_and_ownership": {"credit_ownership_statement": "Not specified"},
        "sustainable_development": {},
    }
    for key, default in defaults.items():
        if key not in data or data[key] is None:
            data[key] = default
        elif isinstance(default, dict):
            for sub_key, sub_val in default.items():
                section = data[key]
                if isinstance(section, dict) and sub_key not in section:
                    section[sub_key] = sub_val

    if "technology" in data and isinstance(data["technology"], dict):
        tech = data["technology"]
        if "waste_type" not in tech or not tech["waste_type"]:
            tech["waste_type"] = ["municipal_solid_waste"]
        if "annual_waste_throughput" not in tech or tech["annual_waste_throughput"] is None:
            tech["annual_waste_throughput"] = 0.1
        if "installed_capacity_mw" not in tech or tech["installed_capacity_mw"] is None:
            tech["installed_capacity_mw"] = 0.0
        if "methodology_ids" not in tech or not tech["methodology_ids"]:
            tech["methodology_ids"] = ["UNKNOWN"]
        if "technology_type" not in tech or tech["technology_type"] is None:
            tech["technology_type"] = "other"

    required_strings: dict[str, dict[str, str]] = {
        "project": {
            "proponent_contact_email": "unknown@unknown.com",
        },
        "location": {
            "region": "Unknown",
        },
        "monitoring": {
            "data_management": "Not specified",
        },
        "safeguards": {
            "no_net_harm_statement": "Not specified",
        },
        "compliance_and_ownership": {
            "credit_ownership_statement": "Not specified",
        },
    }
    for section_key, fields in required_strings.items():
        if section_key in data and isinstance(data[section_key], dict):
            for field, default in fields.items():
                if data[section_key].get(field) is None:
                    data[section_key][field] = default

    return data


def extract_project_input(
    source: str | Path,
    provider: BaseProvider,
    *,
    max_chars: int = _MAX_DOC_CHARS,
) -> ProjectInput:
    """Extract a ProjectInput from a document or raw text string.

    Args:
        source: File path (DOCX/PDF/text) or raw text string.
        provider: LLM provider to use for extraction.
        max_chars: Max characters to send to the LLM.

    Returns:
        Validated ProjectInput with extraction_provenance populated.

    Raises:
        ExtractionError: If extraction or parsing fails.
    """
    source_path: str | None = None

    if isinstance(source, Path) or (isinstance(source, str) and Path(source).exists()):
        path = Path(source)
        if not path.exists():
            raise ExtractionError(f"File not found: {path}")
        source_path = str(path)
        document_text = _read_text_from_file(path)
        logger.info("extract_text_loaded", path=source_path, chars=len(document_text))
    else:
        document_text = str(source)
        logger.info("extract_raw_text", chars=len(document_text))

    if not document_text.strip():
        raise ExtractionError("Document is empty or contains no extractable text")

    prompt = _build_extraction_prompt(document_text[:max_chars])

    response = provider.draft_section(
        section_id="extract",
        sub_section_id="project_input",
        prompt=prompt,
        provenance=["[EXTRACTION: document intake]"],
        max_chars=16000,
    )

    if response.confidence == "UNSUPPORTED" and "ERROR" in response.text:
        raise ExtractionError(f"Provider returned error: {response.text[:500]}")

    raw_data = _parse_yaml_from_response(response.text)
    cleaned_data, missing_fields = _clean_missing_markers(raw_data)
    model_name = getattr(provider, "_model", None)
    if model_name is not None and not isinstance(model_name, str):
        model_name = None
    provenance = _build_provenance(cleaned_data, missing_fields, source_path, model_name)

    cleaned_data = _apply_schema_defaults(cleaned_data)

    try:
        project_input = ProjectInput.model_validate(cleaned_data)
    except Exception as exc:
        raise ExtractionError(f"Extracted data failed ProjectInput validation: {exc}") from exc

    project_input.extraction_provenance = provenance
    return project_input
