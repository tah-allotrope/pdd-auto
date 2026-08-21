"""Tests for the S-1 Markdown block scanner (2026-08-21 real-output-fidelity plan)."""

from __future__ import annotations

from pdd_agent.export.markdown_docx import (
    _split_inline_runs,
    clean_math_text,
    parse_pipe_table,
    render_markdown_body,
)


def _doc():
    from docx import Document

    return Document()


def _paragraph_texts(doc):
    return [p.text for p in doc.paragraphs]


class TestParsePipeTable:
    def test_basic_table(self):
        assert parse_pipe_table(["| A | B |", "|---|---|", "| 1 | 2 |"]) == [
            ["A", "B"],
            ["1", "2"],
        ]

    def test_no_alignment_row_is_not_a_table(self):
        assert parse_pipe_table(["| A | B |", "| 1 | 2 |"]) is None

    def test_ragged_row_padded(self):
        assert parse_pipe_table(["| A | B | C |", "|---|---:|---|", "| 1 | 2 |"]) == [
            ["A", "B", "C"],
            ["1", "2", ""],
        ]

    def test_overflow_truncated(self):
        assert parse_pipe_table(["| A |", "|---|", "| 1 | 2 | 3 |"]) == [["A"], ["1"]]

    def test_single_line_is_not_a_table(self):
        assert parse_pipe_table(["| A | B |"]) is None


class TestCleanMathText:
    def test_display_math_cleanup(self):
        assert (
            clean_math_text("$$BE_y = \\sum_t \\left( X \\right) \\times 2$$")
            == "BE_y = Σ_t ( X ) × 2"
        )

    def test_inline_subscript_cleanup(self):
        assert clean_math_text("$BE_{CH4,t,y}$") == "BE_CH4,t,y"

    def test_unknown_command_keeps_letters(self):
        assert clean_math_text("$\\phi_{y}$") == "phi_y"


class TestSplitInlineRuns:
    def test_bold_span(self):
        assert _split_inline_runs("Total is **487,710.99** tCO2e") == [
            ("plain", "Total is "),
            ("bold", "487,710.99"),
            ("plain", " tCO2e"),
        ]

    def test_plain_text(self):
        assert _split_inline_runs("plain text") == [("plain", "plain text")]

    def test_code_and_italic(self):
        assert _split_inline_runs("`code` and *em") == [
            ("code", "code"),
            ("plain", " and "),
            ("italic", "em"),
        ]

    def test_unpaired_double_delimiter_left_alone(self):
        assert _split_inline_runs("unbalanced ** marker") == [("plain", "unbalanced ** marker")]

    def test_inline_math_becomes_cleaned_italic(self):
        style, text = _split_inline_runs("value $x_{y}$ end")[1]
        assert style == "italic"
        assert text == "x_y"


class TestRenderMarkdownBody:
    def test_atx_heading_becomes_word_heading(self):
        doc = _doc()
        render_markdown_body(doc, "# Baseline Emissions")
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 1
        assert headings[0].style.name.startswith("Heading 3")
        assert headings[0].text == "Baseline Emissions"
        assert not any("#" in t for t in _paragraph_texts(doc))

    def test_deep_heading_capped_at_level_four(self):
        doc = _doc()
        render_markdown_body(doc, "### Deep")
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings[0].style.name.startswith("Heading 4")

    def test_pipe_table_becomes_word_table(self):
        doc = _doc()
        render_markdown_body(doc, "| A | B |\n|---|---|\n| 1 | 2 |")
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "A"
        assert all("|" not in t for t in _paragraph_texts(doc))

    def test_bullet_list(self):
        doc = _doc()
        render_markdown_body(doc, "- one\n- two")
        bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert [p.text for p in bullets] == ["one", "two"]

    def test_numbered_list(self):
        doc = _doc()
        render_markdown_body(doc, "1. first\n2. second")
        numbers = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert [p.text for p in numbers] == ["first", "second"]

    def test_fenced_code_is_monospace_and_not_a_table(self):
        doc = _doc()
        render_markdown_body(doc, "```\nraw | text\n```")
        paragraphs = [p for p in doc.paragraphs if p.text == "raw | text"]
        assert len(paragraphs) == 1
        assert paragraphs[0].runs[0].font.name == "Consolas"
        assert len(doc.tables) == 0

    def test_display_math_collected_and_rendered_centred_italic(self):
        doc = _doc()
        collected: list[str] = []
        render_markdown_body(doc, "$$X = 1$$", on_display_math=collected.append)
        assert collected == ["X = 1"]
        paragraph = [p for p in doc.paragraphs if p.text == "X = 1"][0]
        assert paragraph.runs[0].italic
        assert len(paragraph.runs) == 1

    def test_empty_text_renders_nothing(self):
        doc = _doc()
        before = len(doc.paragraphs)
        render_markdown_body(doc, "")
        assert len(doc.paragraphs) == before

    def test_broken_table_falls_through_to_paragraphs(self):
        doc = _doc()
        render_markdown_body(doc, "| broken\n| still broken")
        texts = _paragraph_texts(doc)
        assert "| broken" in texts
        assert "| still broken" in texts
        assert len(doc.tables) == 0

    def test_inline_emphasis_markers_absent(self):
        doc = _doc()
        render_markdown_body(doc, "**bold** and *soft* words")
        joined = "\n".join(_paragraph_texts(doc))
        assert "**" not in joined
        assert "bold" in joined
