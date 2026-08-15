"""Tests for VCS v4.4 table renderers introduced in Phase 01."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path


from docx import Document
from pdd_agent.export.docx_export import (
    render_cover_metadata_table,
    render_audit_history_table,
    render_proponent_table,
    render_ghg_boundary_table,
    render_applicability_table,
    render_monitoring_fixed_params_table,
    render_monitoring_tracked_params_table,
    render_risk_assessment_table,
    render_emissions_summary_table,
    render_sustainable_development_table,
    render_data_gaps_table,
    render_tbd_appendix,
    export_run_to_docx,
)


def _docx_xml(doc) -> str:
    from io import BytesIO

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    with zipfile.ZipFile(buffer) as archive:
        return archive.read("word/document.xml").decode("utf-8")


class TestTableRenderers:
    """Verify each VCS table renderer produces correct headers and row counts."""

    def test_cover_metadata(self, tmp_path: Path):
        doc = Document()
        data = {
            "project_title": "Inegol WTE Project",
            "project_id": "VCS-1234",
            "crediting_period": "01-Jan-2024 to 31-Dec-2030",
            "original_issue_date": "15-March-2024",
            "most_recent_issue_date": "20-May-2026",
            "version": "Draft 0.1",
            "vcs_standard_version": "v4.4",
            "prepared_by": "Allotrope VC",
        }
        table = render_cover_metadata_table(doc, data)
        assert table is not None
        assert len(table.rows) == 8
        xml = _docx_xml(doc)
        assert "Project title" in xml
        assert "Inegol WTE Project" in xml
        assert "VCS Standard Version" in xml

    def test_audit_history(self, tmp_path: Path):
        doc = Document()
        data = {
            "audits": [
                {
                    "audit_type": "validation",
                    "period": "2024-2030",
                    "program": "VCS",
                    "vvb_name": "TUV SUD",
                    "number_of_years": 7,
                },
                {
                    "audit_type": "verification",
                    "period": "2024",
                    "program": "VCS",
                    "vvb_name": "TUV SUD",
                    "number_of_years": 1,
                },
            ]
        }
        table = render_audit_history_table(doc, data)
        assert table is not None
        assert len(table.rows) == 3
        xml = _docx_xml(doc)
        assert "Audit type" in xml
        assert "Validation/verification body name" in xml
        assert "TUV SUD" in xml

    def test_proponent(self, tmp_path: Path):
        doc = Document()
        data = {
            "org_name": "Allotrope",
            "contact_name": "A. B.",
            "title": "Director",
            "address": "Istanbul",
            "telephone": "+90",
            "email": "a@b.com",
        }
        table = render_proponent_table(doc, data)
        assert table is not None
        assert len(table.rows) == 6
        xml = _docx_xml(doc)
        assert "Organization name" in xml
        assert "Allotrope" in xml

    def test_ghg_boundary(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {
                    "scenario": "Baseline",
                    "source": "SWDS",
                    "gas": "CH4",
                    "included": "Yes",
                    "justification": "Major source",
                },
                {
                    "scenario": "Project",
                    "source": "Engines",
                    "gas": "CO2",
                    "included": "Yes",
                    "justification": "Displacement",
                },
            ]
        }
        table = render_ghg_boundary_table(doc, data)
        assert table is not None
        assert len(table.rows) == 3
        assert len(table.columns) == 5
        xml = _docx_xml(doc)
        assert "Scenario" in xml
        assert "Justification" in xml

    def test_applicability(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {
                    "methodology": "ACM0022",
                    "condition": "Fresh waste treatment",
                    "justification": "Compliant",
                },
                {
                    "methodology": "TOOL02",
                    "condition": "Baseline scenario",
                    "justification": "Used",
                },
            ]
        }
        table = render_applicability_table(doc, data)
        assert table is not None
        assert len(table.rows) == 3
        assert len(table.columns) == 3
        xml = _docx_xml(doc)
        assert "Methodology/tool" in xml
        assert "Compliant" in xml

    def test_monitoring_fixed_params(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {
                    "parameter": "Waste throughput",
                    "unit": "tonnes",
                    "description": "Annual waste",
                    "value": "100,000",
                    "source": "Weighbridge",
                    "comments": "Calibrated",
                },
            ]
        }
        table = render_monitoring_fixed_params_table(doc, data)
        assert table is not None
        assert len(table.rows) == 2
        assert len(table.columns) == 6
        xml = _docx_xml(doc)
        assert "Data/parameter" in xml
        assert "Weighbridge" in xml

    def test_monitoring_tracked_params(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {
                    "parameter": "Electricity export",
                    "unit": "MWh",
                    "description": "Net export",
                    "frequency": "Monthly",
                    "equipment": "Revenue meter",
                    "qa_qc": "Annual calibration",
                },
            ]
        }
        table = render_monitoring_tracked_params_table(doc, data)
        assert table is not None
        assert len(table.rows) == 2
        assert len(table.columns) == 6
        xml = _docx_xml(doc)
        assert "Frequency" in xml
        assert "QA/QC" in xml

    def test_risk_assessment(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {"category": "Environmental", "risks": "Leachate", "mitigation": "Treatment plant"},
                {
                    "category": "Health & Safety",
                    "risks": "Biogas exposure",
                    "mitigation": "Gas detection",
                },
            ]
        }
        table = render_risk_assessment_table(doc, data)
        assert table is not None
        assert len(table.rows) == 3
        assert len(table.columns) == 3
        xml = _docx_xml(doc)
        assert "Risk category" in xml
        assert "Biogas exposure" in xml

    def test_emissions_summary(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {"period": "2024", "value": "75,000"},
                {"period": "2025", "value": "75,000"},
            ],
            "total": "750,000",
        }
        table = render_emissions_summary_table(doc, data)
        assert table is not None
        assert len(table.rows) == 4
        assert len(table.columns) == 2
        xml = _docx_xml(doc)
        assert "Calendar year of crediting period" in xml
        assert "750,000" in xml

    def test_sustainable_development(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {
                    "area": "Climate mitigation",
                    "contribution": "Avoided methane",
                    "monitoring": "Annual calculations",
                },
                {
                    "area": "Clean energy",
                    "contribution": "Renewable electricity",
                    "monitoring": "Metered export",
                },
            ]
        }
        table = render_sustainable_development_table(doc, data)
        assert table is not None
        assert len(table.rows) == 3
        assert len(table.columns) == 3
        xml = _docx_xml(doc)
        assert "Sustainable development area" in xml
        assert "Avoided methane" in xml

    def test_data_gaps(self, tmp_path: Path):
        doc = Document()
        data = {
            "entries": [
                {
                    "topic": "Stakeholder consultation",
                    "gap": "No records provided",
                    "evidence": "Meeting minutes",
                },
            ]
        }
        table = render_data_gaps_table(doc, data)
        assert table is not None
        assert len(table.rows) == 2
        assert len(table.columns) == 3
        xml = _docx_xml(doc)
        assert "Topic" in xml
        assert "Needed evidence" in xml

    def test_tbd_appendix(self, tmp_path: Path):
        doc = Document()
        report = {
            "items": [
                {
                    "section": "1.10",
                    "marker": "[TBD]",
                    "context": "Missing emission factor",
                    "suggested_evidence": "Grid authority letter",
                },
            ]
        }
        table = render_tbd_appendix(doc, report)
        assert table is not None
        assert len(table.rows) == 2
        assert len(table.columns) == 4
        xml = _docx_xml(doc)
        assert "Appendix C - Data Gaps and Evidence Requirements" in xml
        assert "Missing emission factor" in xml

    def test_tbd_appendix_empty(self, tmp_path: Path):
        doc = Document()
        report = {"items": []}
        table = render_tbd_appendix(doc, report)
        assert table is None
        xml = _docx_xml(doc)
        assert "Appendix C" in xml
        assert "No TBD markers were detected" in xml


class TestCalcStructuredContentEndToEnd:
    """Phase 05: structured_content carried on a run's section renders as a live table."""

    def test_emissions_summary_structured_content_renders_table(self, tmp_path: Path, monkeypatch):
        run_dir = tmp_path / "data" / "runs"
        run_dir.mkdir(parents=True, exist_ok=True)
        payload = {
            "run_id": "structured-content-run",
            "project_name": "Soc Son waste to power plant project",
            "provider": "noop",
            "assumption_register": {"assumptions": [], "guardrails": {}},
            "sections": [
                {
                    "section_id": "4",
                    "sub_section_id": "4.4",
                    "text": "Narrative that is replaced by the calc-driven table.",
                    "confidence": "HIGH",
                    "provenance": [],
                    "issues": [],
                    "provider": "noop",
                    "structured_content": {
                        "table_type": "emissions_summary",
                        "data": {
                            "entries": [
                                {"period": year, "value": f"{75_000 * year:,.0f}"}
                                for year in range(1, 8)
                            ],
                            "total": "1,837,500",
                        },
                    },
                }
            ],
            "notes": [],
        }
        run_path = run_dir / "structured-content-run.json"
        run_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        monkeypatch.setattr("pdd_agent.export.docx_export._DRAFT_RUNS_DIR", run_dir)

        output = export_run_to_docx(
            "structured-content-run", output_path=tmp_path / "structured.docx"
        )
        with zipfile.ZipFile(output) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")

        assert "Calendar year of crediting period" in xml
        assert "Estimated GHG emission reductions or removals" in xml
        assert "1,837,500" in xml
        # PHASE-03 fix: structured_content used to delete the section's prose
        # (old contract asserted `"Narrative that is replaced" not in xml`).
        # Prose and table now render together.
        assert "Narrative that is replaced" in xml

    def test_nonexistent_table_type_renders_prose_only_and_does_not_raise(
        self, tmp_path: Path, monkeypatch
    ):
        """Section 3.3 test spec: an unresolved table_type still renders the
        section's prose and adds no table; export must not raise."""
        run_dir = tmp_path / "data" / "runs"
        run_dir.mkdir(parents=True, exist_ok=True)
        payload = {
            "run_id": "unresolved-table-type-run",
            "project_name": "Soc Son waste to power plant project",
            "provider": "noop",
            "assumption_register": {"assumptions": [], "guardrails": {}},
            "sections": [
                {
                    "section_id": "3",
                    "sub_section_id": "3.3",
                    "text": "Narrative that must survive an unresolved table type.",
                    "confidence": "HIGH",
                    "provenance": [],
                    "issues": [],
                    "provider": "noop",
                    "structured_content": {"table_type": "nonexistent_type", "data": {}},
                }
            ],
            "notes": [],
        }
        run_path = run_dir / "unresolved-table-type-run.json"
        run_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        monkeypatch.setattr("pdd_agent.export.docx_export._DRAFT_RUNS_DIR", run_dir)

        output = export_run_to_docx(
            "unresolved-table-type-run", output_path=tmp_path / "unresolved.docx"
        )
        with zipfile.ZipFile(output) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")

        assert "Narrative that must survive an unresolved table type." in xml

    def test_section_without_structured_content_renders_prose_only(
        self, tmp_path: Path, monkeypatch
    ):
        """Unchanged behaviour: a section with no structured_content renders
        just its prose paragraph, as before this phase."""
        run_dir = tmp_path / "data" / "runs"
        run_dir.mkdir(parents=True, exist_ok=True)
        payload = {
            "run_id": "no-structured-content-run",
            "project_name": "Soc Son waste to power plant project",
            "provider": "noop",
            "assumption_register": {"assumptions": [], "guardrails": {}},
            "sections": [
                {
                    "section_id": "2",
                    "sub_section_id": "2.1",
                    "text": "No net harm.",
                    "confidence": "HIGH",
                    "provenance": [],
                    "issues": [],
                    "provider": "noop",
                }
            ],
            "notes": [],
        }
        run_path = run_dir / "no-structured-content-run.json"
        run_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        monkeypatch.setattr("pdd_agent.export.docx_export._DRAFT_RUNS_DIR", run_dir)

        output = export_run_to_docx(
            "no-structured-content-run", output_path=tmp_path / "no-structured.docx"
        )
        with zipfile.ZipFile(output) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")

        assert "No net harm." in xml
